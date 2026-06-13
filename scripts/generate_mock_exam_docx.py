#!/usr/bin/env python3
"""Generate a standardized mock exam DOCX in landscape A4 format.

This script creates a professional exam paper that looks like a real university
exam, with proper formatting, headers, sections, and an answer key.

Usage:
    python scripts/generate_mock_exam_docx.py \
        --input examples/mock_exam.sample.json \
        --output mock_exam.docx

Dependency:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches


COLORS = {
    "primary_blue": "1F5D8C",
    "heading_navy": "1E3A56",
    "light_blue": "EAF3FB",
    "light_yellow": "FFF6D9",
    "light_red": "FDEBEC",
    "light_gray": "F4F6F8",
    "border_gray": "D8DEE6",
    "white": "FFFFFF",
    "dark_text": "333333",
}

DEFAULT_FONT_EN = "Aptos"
DEFAULT_FONT_CJK = "Microsoft YaHei"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLORS["border_gray"], size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "left", "bottom", "right"):
            element = p_bdr.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                p_bdr.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:space"), "4")
            element.set(qn("w:color"), border)


def set_keep_with_next(paragraph) -> None:
    """Set paragraph to keep with next paragraph (prevent page break between them)."""
    pPr = paragraph._p.get_or_add_pPr()
    keepNext = pPr.find(qn("w:keepNext"))
    if keepNext is None:
        keepNext = OxmlElement("w:keepNext")
        pPr.append(keepNext)


def set_keep_lines_together(paragraph) -> None:
    """Set paragraph to keep all lines together (prevent page break within paragraph)."""
    pPr = paragraph._p.get_or_add_pPr()
    keepLines = pPr.find(qn("w:keepLines"))
    if keepLines is None:
        keepLines = OxmlElement("w:keepLines")
        pPr.append(keepLines)


def set_page_break_before(paragraph) -> None:
    """Force page break before this paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pageBreakBefore = pPr.find(qn("w:pageBreakBefore"))
    if pageBreakBefore is None:
        pageBreakBefore = OxmlElement("w:pageBreakBefore")
        pPr.append(pageBreakBefore)


def set_run_font(run, size: int | float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None,
                 underline: bool | None = None) -> None:
    run.font.name = DEFAULT_FONT_EN
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), DEFAULT_FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.font.underline = underline
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(document: Document, text: str = "", style: str | None = None, *,
                  bold: bool = False, size: float | None = None,
                  color: str | None = None, align=None,
                  keep_with_next: bool = False) -> Any:
    paragraph = document.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size or 10.5, bold=bold, color=color or COLORS["dark_text"])
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    if align is not None:
        paragraph.alignment = align
    if keep_with_next:
        set_keep_with_next(paragraph)
    return paragraph


def add_heading(document: Document, text: str, level: int = 1) -> Any:
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=14, bold=True, color=COLORS["primary_blue"])
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=COLORS["heading_navy"])
    else:
        set_run_font(run, size=11, bold=True, color=COLORS["heading_navy"])
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    set_keep_with_next(paragraph)
    return paragraph


def add_callout(document: Document, title: str, text: str, kind: str = "main") -> None:
    fill_map = {
        "main": COLORS["light_blue"],
        "notice": COLORS["light_yellow"],
        "mistake": COLORS["light_red"],
        "example": COLORS["light_gray"],
    }
    label_color = COLORS["primary_blue"] if kind == "main" else COLORS["heading_navy"]
    fill = fill_map.get(kind, COLORS["light_gray"])

    p = document.add_paragraph()
    shade_paragraph(p, fill, COLORS["border_gray"])
    label = p.add_run(f"{title}: ")
    set_run_font(label, size=10, bold=True, color=label_color)
    body = p.add_run(text)
    set_run_font(body, size=10)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    set_keep_with_next(p)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for i, item in enumerate(items):
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(str(item))
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)
        if i < len(items) - 1:
            set_keep_with_next(p)


def add_table(document: Document, headers: list[str], rows: list[list[Any]],
              col_widths: list[float] | None = None) -> None:
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # Header row
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLORS["primary_blue"])
        set_cell_border(cell, COLORS["primary_blue"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(header))
        set_run_font(run, size=9.5, bold=True, color=COLORS["white"])

    # Data rows
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row[:len(headers)]):
            cell = cells[col_index]
            if row_index % 2 == 1:
                set_cell_shading(cell, "F9FBFD")
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.5)
    document.add_paragraph()


def add_separator(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLORS["border_gray"])
    pBdr.append(bottom)
    pPr.append(pBdr)


def configure_page_layout(document: Document, metadata: dict[str, Any]) -> None:
    section = document.sections[0]

    # Orientation from JSON: "landscape" or "portrait"
    orientation = metadata.get("orientation", "landscape").lower()
    if orientation == "portrait":
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    else:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)

    # Margins
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Header
    header = section.header.paragraphs[0]
    header.text = metadata.get("header", metadata.get("course", "期末模拟试卷"))
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, size=9, color=COLORS["heading_navy"])

    # Footer
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(metadata.get("footer", "本试卷仅供学习参考，不作为正式考试依据"))
    set_run_font(run, size=8, color="999999")

    # Default font
    styles = document.styles
    styles["Normal"].font.name = DEFAULT_FONT_EN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT_CJK)
    styles["Normal"].font.size = Pt(10.5)


def add_exam_header(document: Document, metadata: dict[str, Any]) -> None:
    # Course name
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metadata.get("course", "课程名称"))
    set_run_font(run, size=18, bold=True, color=COLORS["primary_blue"])
    set_keep_with_next(p)

    # Exam title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metadata.get("exam_title", "期末模拟试卷"))
    set_run_font(run, size=16, bold=True, color=COLORS["heading_navy"])
    set_keep_with_next(p)

    # Info line
    info_parts = []
    if metadata.get("duration"):
        info_parts.append(f"考试时间：{metadata['duration']}")
    if metadata.get("total_marks"):
        info_parts.append(f"满分：{metadata['total_marks']} 分")
    if metadata.get("mode"):
        info_parts.append(f"考试形式：{metadata['mode']}")

    if info_parts:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("　　".join(info_parts))
        set_run_font(run, size=11, bold=True, color=COLORS["dark_text"])
        set_keep_with_next(p)

    # Separator line
    add_separator(document)

    # Instructions
    instructions = metadata.get("instructions", [
        "本试卷共 X 大题，满分 100 分，请在规定时间内完成。",
        "答题前请仔细阅读各题要求，合理分配时间。",
        "书写工整，步骤清晰，未注明解题过程的计算题可不给分。",
    ])
    add_callout(document, "注意事项", "；".join(instructions), "notice")

    add_separator(document)


def add_question_block(document: Document, question: dict[str, Any]) -> None:
    q_num = question.get("number", "")
    q_type = question.get("type", "")
    q_text = question.get("text", "")
    q_marks = question.get("marks", "")
    q_time = question.get("time", "")
    q_options = question.get("options", [])
    q_sub = question.get("sub_questions", [])
    q_page_break_before = question.get("page_break_before", False)

    # Optional page break before this question
    if q_page_break_before:
        p = document.add_paragraph()
        set_page_break_before(p)

    # Question header
    header_parts = [f"{q_num}."]
    if q_type:
        header_parts.append(f"【{q_type}】")
    if q_marks:
        header_parts.append(f"（{q_marks} 分）")
    if q_time:
        header_parts.append(f"（约 {q_time}）")

    p = document.add_paragraph()
    run = p.add_run(" ".join(header_parts))
    set_run_font(run, size=11, bold=True, color=COLORS["heading_navy"])
    set_keep_with_next(p)

    # Question text
    if q_text:
        p = document.add_paragraph()
        run = p.add_run(q_text)
        set_run_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(4)
        set_keep_with_next(p)

    # Options (for MCQ)
    for i, option in enumerate(q_options):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(option)
        set_run_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(1)
        if i < len(q_options) - 1:
            set_keep_with_next(p)

    # Sub-questions
    for i, sub in enumerate(q_sub):
        sub_num = sub.get("number", "")
        sub_text = sub.get("text", "")
        sub_marks = sub.get("marks", "")

        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        header = f"{sub_num}. "
        if sub_marks:
            header += f"（{sub_marks} 分）"
        run = p.add_run(header + sub_text)
        set_run_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(2)
        if i < len(q_sub) - 1:
            set_keep_with_next(p)

    # Answer space indicator
    if question.get("answer_lines"):
        for _ in range(question["answer_lines"]):
            p = document.add_paragraph()
            run = p.add_run("　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　")
            set_run_font(run, size=10, color="CCCCCC")
            p.paragraph_format.space_after = Pt(1)


def add_section_header(document: Document, section: dict[str, Any], is_first: bool = False) -> None:
    title = section.get("title", "")
    desc = section.get("description", "")
    marks = section.get("total_marks", "")
    page_break = section.get("page_break_before", True)  # Default: page break before each section

    # Page break before section (except the first one)
    if page_break and not is_first:
        p = document.add_paragraph()
        set_page_break_before(p)

    p = document.add_paragraph()
    shade_paragraph(p, COLORS["light_blue"])

    title_text = title
    if marks:
        title_text += f"（共 {marks} 分）"
    run = p.add_run(title_text)
    set_run_font(run, size=12, bold=True, color=COLORS["primary_blue"])
    set_keep_with_next(p)

    if desc:
        p = document.add_paragraph()
        run = p.add_run(desc)
        set_run_font(run, size=10, italic=True, color="666666")
        set_keep_with_next(p)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_answer_section(document: Document, exam_data: dict[str, Any]) -> None:
    # Page break before answer section
    p = document.add_paragraph()
    set_page_break_before(p)

    # Answer section title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("参考答案与评分标准")
    set_run_font(run, size=16, bold=True, color=COLORS["primary_blue"])
    set_keep_with_next(p)

    # Course info
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"课程名称：{exam_data.get('metadata', {}).get('course', '')}")
    set_run_font(run, size=11, color=COLORS["heading_navy"])
    set_keep_with_next(p)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"试卷版本：{exam_data.get('metadata', {}).get('version', 'A')}")
    set_run_font(run, size=11, color=COLORS["heading_navy"])

    add_separator(document)

    # Answers
    for i, answer in enumerate(exam_data.get("answers", [])):
        q_num = answer.get("number", "")
        answer_text = answer.get("answer", "")
        explanation = answer.get("explanation", "")
        scoring = answer.get("scoring", [])
        common_errors = answer.get("common_errors", [])

        # Answer header
        p = document.add_paragraph()
        run = p.add_run(f"第 {q_num} 题")
        set_run_font(run, size=11, bold=True, color=COLORS["heading_navy"])
        set_keep_with_next(p)

        # Answer text
        if answer_text:
            p = document.add_paragraph()
            shade_paragraph(p, COLORS["light_gray"])
            run = p.add_run(f"【答案】{answer_text}")
            set_run_font(run, size=10.5)
            set_keep_with_next(p)

        # Explanation
        if explanation:
            p = document.add_paragraph()
            run = p.add_run(f"【解析】{explanation}")
            set_run_font(run, size=10)
            set_keep_with_next(p)

        # Scoring rubric
        if scoring:
            add_table(document,
                     ["评分项", "分值", "给分条件"],
                     scoring)

        # Common errors
        if common_errors:
            for error in common_errors:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(f"✗ {error}")
                set_run_font(run, size=9.5, color="CC4444")

        add_separator(document)


def add_post_exam_diagnosis(document: Document, exam_data: dict[str, Any]) -> None:
    diagnosis = exam_data.get("diagnosis", {})
    if not diagnosis:
        return

    # Page break before diagnosis
    p = document.add_paragraph()
    set_page_break_before(p)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("考后诊断分析")
    set_run_font(run, size=14, bold=True, color=COLORS["primary_blue"])

    add_separator(document)

    # By chapter
    if diagnosis.get("by_chapter"):
        add_heading(document, "按章节诊断", 2)
        add_table(document,
                 ["章节", "涉及题号", "分值", "掌握程度", "复习建议"],
                 diagnosis["by_chapter"])

    # By cognitive level
    if diagnosis.get("by_level"):
        add_heading(document, "按认知层级诊断", 2)
        add_table(document,
                 ["层级", "涉及题号", "分值", "复习建议"],
                 diagnosis["by_level"])

    # Weak points
    if diagnosis.get("weak_points"):
        add_heading(document, "薄弱点修复建议", 2)
        add_bullet_list(document, diagnosis["weak_points"])


def build_exam_docx(data: dict[str, Any], output_path: Path) -> None:
    document = Document()
    metadata = data.get("metadata", {})

    configure_page_layout(document, metadata)
    add_exam_header(document, metadata)

    # Add sections and questions
    sections = data.get("sections", [])
    for idx, section in enumerate(sections):
        add_section_header(document, section, is_first=(idx == 0))
        for question in section.get("questions", []):
            add_question_block(document, question)
            add_separator(document)

    # Add answer key
    add_answer_section(document, data)

    # Add post-exam diagnosis
    add_post_exam_diagnosis(document, data)

    document.save(output_path)
    print(f"Created {output_path}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a mock exam DOCX (landscape A4).")
    parser.add_argument("--input", required=True, help="Path to mock exam JSON.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    data = json.loads(input_path.read_text(encoding="utf-8"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    build_exam_docx(data, output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
