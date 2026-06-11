#!/usr/bin/env python3
"""Generate a polished university final-review DOCX from structured JSON.

This script is intentionally part of the skill's core capability: it turns
review content into a styled Word handout using the visual rules documented in
`docs/en/docx-style-guide.md` and `docs/zh-CN/docx-style-guide.md`.

Usage:
    python scripts/generate_styled_docx.py \
        --input examples/review_content.sample.json \
        --output final_review.docx

Dependency:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


COLORS = {
    "primary_blue": "1F5D8C",
    "heading_navy": "1E3A56",
    "light_blue": "EAF3FB",
    "light_yellow": "FFF6D9",
    "light_red": "FDEBEC",
    "light_gray": "F4F6F8",
    "border_gray": "D8DEE6",
    "white": "FFFFFF",
}

DEFAULT_FONT_EN = "Aptos"
DEFAULT_FONT_CJK = "Microsoft YaHei"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLORS["border_gray"], size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)

    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "left", "bottom", "right"):
            element = p_bdr.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                p_bdr.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "4")
            element.set(qn("w:color"), border)


def set_run_font(run, size: int | float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = DEFAULT_FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(document: Document, text: str = "", style: str | None = None, *, bold: bool = False) -> Any:
    paragraph = document.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.18
    return paragraph


def add_heading(document: Document, text: str, level: int = 1) -> Any:
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=COLORS["primary_blue"])
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=COLORS["heading_navy"])
    else:
        set_run_font(run, size=11.5, bold=True, color=COLORS["heading_navy"])
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_callout(document: Document, title: str, text: str | Iterable[str], kind: str = "main") -> None:
    fill_map = {
        "main": COLORS["light_blue"],
        "notice": COLORS["light_yellow"],
        "mistake": COLORS["light_red"],
        "example": COLORS["light_gray"],
    }
    label_color = COLORS["primary_blue"] if kind == "main" else COLORS["heading_navy"]
    fill = fill_map.get(kind, COLORS["light_gray"])

    p = document.add_paragraph()
    shade_paragraph(p, fill, COLORS["border_gray"])
    label = p.add_run(f"{title}: ")
    set_run_font(label, bold=True, color=label_color)
    if isinstance(text, str):
        body = p.add_run(text)
        set_run_font(body)
    else:
        body = p.add_run("; ".join(str(item) for item in text))
        set_run_font(body)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(str(item))
        set_run_font(run)
        p.paragraph_format.space_after = Pt(2)


def add_table(document: Document, headers: list[str], rows: list[dict[str, Any] | list[Any]]) -> None:
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        set_cell_shading(cell, COLORS["primary_blue"])
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(header))
        set_run_font(run, bold=True, color=COLORS["white"])

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        values = [row.get(header, "") for header in headers] if isinstance(row, dict) else row
        for col_index, value in enumerate(values[: len(headers)]):
            cell = cells[col_index]
            if row_index % 2 == 1:
                set_cell_shading(cell, "F9FBFD")
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.5)
    document.add_paragraph()


def configure_document(document: Document, metadata: dict[str, Any]) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    header = section.header.paragraphs[0]
    header.text = metadata.get("header", metadata.get("course", "University Final Review"))
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, size=9, color=COLORS["heading_navy"])

    footer = section.footer.paragraphs[0]
    footer.text = metadata.get("footer", "Review Notes and Exam Focus")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color="666666")

    styles = document.styles
    styles["Normal"].font.name = DEFAULT_FONT_EN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT_CJK)
    styles["Normal"].font.size = Pt(10.5)


def add_cover(document: Document, metadata: dict[str, Any], chapters: list[dict[str, Any]]) -> None:
    breadcrumb = metadata.get("breadcrumb", metadata.get("course", "Course Review"))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(breadcrumb)
    set_run_font(run, size=9.5, color="666666")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(metadata.get("title", metadata.get("course", "University Final Review")))
    set_run_font(run, size=22, bold=True, color=COLORS["primary_blue"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(metadata.get("subtitle", "Detailed Review Notes and Exam Focus"))
    set_run_font(run, size=13, bold=True, color=COLORS["heading_navy"])

    if metadata.get("main_thread"):
        add_callout(document, metadata.get("main_thread_label", "Main thread"), metadata["main_thread"], "main")

    if metadata.get("objectives"):
        add_heading(document, "Learning Objectives", 2)
        add_bullet_list(document, metadata["objectives"])

    if chapters:
        add_heading(document, "Table of Contents", 2)
        for index, chapter in enumerate(chapters, 1):
            add_paragraph(document, f"{index}. {chapter.get('title', 'Untitled Chapter')}")

    document.add_page_break()


def add_chapter(document: Document, chapter: dict[str, Any], index: int) -> None:
    add_heading(document, f"Chapter {index}: {chapter.get('title', 'Untitled Chapter')}", 1)

    if chapter.get("main_thread"):
        add_callout(document, chapter.get("main_thread_label", "Main thread"), chapter["main_thread"], "main")

    if chapter.get("objectives"):
        add_heading(document, "Learning Objectives", 2)
        add_bullet_list(document, chapter["objectives"])

    if chapter.get("concepts"):
        add_heading(document, "Core Concepts", 2)
        add_table(document, chapter.get("concept_headers", ["Concept", "Meaning", "Example"]), chapter["concepts"])

    for section in chapter.get("sections", []):
        add_heading(document, section.get("title", "Section"), 2)
        for paragraph in section.get("paragraphs", []):
            add_paragraph(document, str(paragraph))

    if chapter.get("formulas"):
        add_heading(document, "Formula / Rule Blocks", 2)
        for item in chapter["formulas"]:
            add_callout(document, item.get("title", "Formula"), item.get("text", ""), "example")

    if chapter.get("examples"):
        add_heading(document, "Worked Examples", 2)
        for item in chapter["examples"]:
            add_callout(document, item.get("title", "Example"), item.get("text", ""), "example")

    for item in chapter.get("notices", []):
        add_callout(document, item.get("title", "Notice"), item.get("text", ""), "notice")

    for item in chapter.get("mistakes", []):
        add_callout(document, item.get("title", "Mistake-prone point"), item.get("text", ""), "mistake")

    if chapter.get("exam_focus"):
        add_heading(document, "Exam Focus", 2)
        add_table(document, chapter.get("exam_focus_headers", ["Item", "Key point", "Exam interpretation"]), chapter["exam_focus"])

    if chapter.get("quick_review"):
        add_heading(document, "One-page Quick Review", 2)
        add_bullet_list(document, chapter["quick_review"])

    if chapter.get("questions"):
        add_heading(document, "Practice Questions", 2)
        for question in chapter["questions"]:
            add_paragraph(document, f"Q: {question.get('question', '')}", bold=True)
            add_paragraph(document, f"A: {question.get('answer', '')}")

    if index != -1:
        document.add_page_break()


def build_docx(data: dict[str, Any], output_path: Path) -> None:
    document = Document()
    metadata = data.get("metadata", {})
    chapters = data.get("chapters", [])
    configure_document(document, metadata)
    add_cover(document, metadata, chapters)
    for index, chapter in enumerate(chapters, 1):
        add_chapter(document, chapter, index)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a styled final-review DOCX from JSON.")
    parser.add_argument("--input", required=True, help="Path to structured review JSON.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    data = json.loads(input_path.read_text(encoding="utf-8"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    build_docx(data, output_path)
    print(f"Created {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
